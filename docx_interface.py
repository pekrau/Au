"DOCX document I/O."

from icecream import ic

import os.path

import docx

import constants
import utils


class Writer:

    def __init__(self, absdirpath, texts):
        self.absdirpath = absdirpath
        self.texts = texts

    def write(self):
        self.document = docx.Document()
        title = os.path.basename(self.absdirpath)
        self.document.add_heading(title, 0)
        prev_dirstack = []
        for filepath in self.texts:
            dirpath = os.path.split(filepath)[0]
            dirstack = utils.split_all(dirpath)
            if dirstack != prev_dirstack:
                if len(dirstack) >= len(prev_dirstack):
                    self.write_section(dirstack)
                prev_dirstack = dirstack
            self.write_text(dirstack, filepath)
        self.document.save(os.path.join(self.absdirpath, title + ".docx"))

    def write_section(self, dirstack):
        if len(dirstack) <= constants.DOCX_PAGEBREAK_LEVEL:
            self.document.add_page_break()
        self.document.add_heading(dirstack[-1], len(dirstack))

    def write_text(self, dirstack, filepath):
        if len(dirstack) < constants.DOCX_PAGEBREAK_LEVEL:
            self.document.add_page_break()
        title = os.path.splitext(os.path.split(filepath)[1])[0]
        self.document.add_heading(title, len(dirstack) + 1)
        parsed = utils.parse(os.path.join(self.absdirpath, filepath))
        self.parse(parsed.ast)

    def parse(self, ast):
        try:
            method = getattr(self, f"parse_{ast['element']}")
        except AttributeError:
            ic("Could not handle ast", ast)
        else:
            method(ast)

    def parse_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.parse(child)

    def parse_paragraph(self, ast):
        self.paragraph = self.document.add_paragraph()
        for child in ast["children"]:
            self.parse(child)

    def parse_raw_text(self, ast):
        line = ast["children"]
        if line[-1] == "\n":
            line[-1] = " "
        self.paragraph.add_run(line)
