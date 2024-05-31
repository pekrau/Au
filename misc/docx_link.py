"""From https://stackoverflow.com/questions/57586400/how-to-create-bookmarks-in-a-word-document-then-create-internal-hyperlinks-to-t

you should change the "r" to "p" in "('//w:r')" :

tag = doc.element.xpath('//w:p')[-1]
"""


from docx import Document
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
    start.set(docx.oxml.ns.qn("w:id"), "0")
    start.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
    end.set(docx.oxml.ns.qn("w:id"), "0")
    end.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(end)


def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")

    # set attribute for link to bookmark
    hyperlink.set(
        docx.oxml.shared.qn("w:anchor"),
        link_to,
    )

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(
            docx.oxml.shared.qn("w:tooltip"),
            tool_tip,
        )

    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True


# test the functions
if __name__ == "__main__":
    # input test document
    doc = Document(r"test_input_1.docx")

    # add a bookmark to every paragraph
    for paranum, paragraph in enumerate(doc.paragraphs):
        add_bookmark(
            paragraph=paragraph,
            bookmark_text=f"{paranum}",
            bookmark_name=f"temp{paranum+1}",
        )

    # add page to the end to put your link
    doc.add_page_break()
    paragraph = doc.add_paragraph("This is where the internal link will live")

    # add a link to the first paragraph
    add_link(
        paragraph=paragraph,
        link_to="temp0",
        text="this is a link to ",
        tool_tip="your message here",
    )

    doc.save(r"output.docx")
