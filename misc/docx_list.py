import docx

document = docx.Document()

paragraph = document.add_paragraph(style="Title")
paragraph.add_run("This is the title.")

paragraph = document.add_paragraph(style="Normal")
paragraph.add_run("Normal text.")

paragraph = document.add_paragraph(style="List Bullet")
paragraph.add_run("List bullet text.")

paragraph = document.add_paragraph(style="List Bullet 2")
paragraph.add_run("List bullet 2 text.")

paragraph = document.add_paragraph(style="List Bullet 3")
paragraph.add_run("List bullet 3 text.")
for pos in range(20):
    paragraph.add_run(" More text.")

paragraph = document.add_paragraph(style="List Continue 3")
paragraph.add_run("Continue List bullet 3 text.")
for pos in range(20):
    paragraph.add_run(" More text.")

paragraph = document.add_paragraph(style="List Bullet")
paragraph.add_run("List bullet text.")
    
paragraph = document.add_paragraph(style="List Number")
paragraph.add_run("List number text.")

paragraph = document.add_paragraph(style="List Number 2")
paragraph.add_run("List number 2 text.")
for pos in range(20):
    paragraph.add_run(" More text.")

paragraph = document.add_paragraph(style="List Number 2")
paragraph.add_run("List number 2 text, next list element.")

paragraph = document.add_paragraph(style="List Number 3")
paragraph.add_run("List number 3 text.")
for pos in range(20):
    paragraph.add_run(" More text.")

paragraph = document.add_paragraph(style="List Number 2")
paragraph.add_run("List number 2 text.")
for pos in range(20):
    paragraph.add_run(" More text.")


document.save("output.docx")
