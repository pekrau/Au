"DOCX export."

from icecream import ic

import copy
import datetime
import os

import tkinter as tk
import tkinter.simpledialog
import tkinter.filedialog

import docx
import docx.oxml

import constants
import utils
from utils import Tr

CODE_STYLE = "Au Code"
CODE_LEFT_INDENT = 30

QUOTE_STYLE = "Au Quote"
QUOTE_LEFT_INDENT = 30
QUOTE_RIGHT_INDENT = 70


class Exporter:
    "DOCX exporter."

    def __init__(self, main, source, config):
        self.main = main
        self.source = source
        self.config = config

    def write(self):
        # Key: canonical; value: dict(id, fullname, ordinal)
        self.indexed = {}
        self.indexed_count = 0
        # Key: reference id; value: dict(id, fullname)
        self.referenced = {}
        self.referenced_count = 0
        # Key: fullname; value: dict(label, ast_children)
        self.footnotes = {}

        self.document = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        if self.main.language:
            styles_element = self.document.styles.element
            rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
            lang_default = rpr_default.xpath("w:lang")[0]
            lang_default.set(docx.oxml.shared.qn("w:val"), self.main.language)

        # Set to A4 page size.
        section = self.document.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)
        section.footer_distance = docx.shared.Mm(12.7)

        # Create style for code.
        style = self.document.styles.add_style(
            CODE_STYLE, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
        )
        style.base_style = self.document.styles["macro"]
        style.paragraph_format.left_indent = docx.shared.Pt(CODE_LEFT_INDENT)
        style.font.name = constants.CODE_FONT

        # Create style for quote.
        style = self.document.styles.add_style(
            QUOTE_STYLE, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
        )
        style.paragraph_format.left_indent = docx.shared.Pt(QUOTE_LEFT_INDENT)
        style.paragraph_format.right_indent = docx.shared.Pt(QUOTE_RIGHT_INDENT)
        style.font.name = constants.QUOTE_FONT

        # Set Dublin core metadata.
        self.document.core_properties.language = self.main.language
        self.document.core_properties.modified = datetime.datetime.now()
        # XXX authors

        self.write_title_page()
        self.write_toc()
        self.write_page_number()
        self.current_text = None
        self.footnote_paragraph = None
        for item in self.source.items:
            if item.is_section:
                self.write_section(item, level=1)
            else:
                self.write_text(item, level=1)
        self.write_references()
        self.write_indexed()

        self.document.save(
            os.path.join(self.config["dirpath"], self.config["filename"])
        )

    def write_title_page(self):
        paragraph = self.document.add_paragraph(style="Title")
        run = paragraph.add_run(self.main.title)
        run.font.size = docx.shared.Pt(constants.FONT_TITLE_SIZE)
        run.font.bold = True

        if self.main.subtitle:
            paragraph = self.document.add_paragraph(style=f"Heading 1")
            paragraph.add_run(self.main.subtitle)

        paragraph.paragraph_format.space_after = docx.shared.Pt(40)
        for author in self.main.authors:
            paragraph = self.document.add_paragraph(style=f"Heading 2")
            paragraph.add_run(author)

        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = docx.shared.Pt(100)

        status = str(
            min([t.status for t in self.source.all_texts] + [max(constants.STATUSES)])
        )
        paragraph.add_run(f"{Tr('Status')}: {Tr(status)}")

        now = datetime.datetime.now().strftime(constants.TIME_ISO_FORMAT)
        self.document.add_paragraph(f"{Tr('Created')}: {now}")

    def write_toc(self):
        # self.document.add_page_break()
        pass

    def write_page_number(self):
        "Display page number in the header."
        # From https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.document.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def write_section(self, section, level):
        if level <= self.config["page_break_level"]:
            self.document.add_page_break()
        self.write_heading(section.heading, level)
        for item in section.items:
            if item.is_section:
                self.write_section(item, level=level + 1)
            else:
                self.write_text(item, level=level + 1)

    def write_text(self, text, level):
        if level <= self.config["page_break_level"]:
            self.document.add_page_break()
        if text.get("display_heading", True):
            self.write_heading(text.heading, level)
        self.list_stack = []
        self.style_stack = ["Normal"]
        self.bold = False
        self.italic = False
        self.current_text = text
        self.render(text.ast)
        self.write_text_footnotes(text)

    def write_heading(self, heading, level):
        level = min(level, constants.MAX_H_LEVEL)
        h = constants.H_LOOKUP[level]
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        paragraph.paragraph_format.left_indent = docx.shared.Pt(h["left_margin"])
        paragraph.paragraph_format.space_after = docx.shared.Pt(h["spacing"])
        run = paragraph.add_run(heading)
        run.font.size = docx.shared.Pt(h["font"][1])

    def write_text_footnotes(self, text):
        "Footnotes at end of the text."
        try:
            footnotes = self.footnotes[text.fullname]
        except KeyError:
            return
        paragraph = self.document.add_heading(Tr("Footnotes"), 6)
        paragraph.paragraph_format.space_before = docx.shared.Pt(25)
        paragraph.paragraph_format.space_after = docx.shared.Pt(10)
        for label in sorted(footnotes.keys()):
            self.footnote_paragraph = self.document.add_paragraph()
            run = self.footnote_paragraph.add_run(f"{label}  ")
            run.italic = True
            for child in footnotes[label]["ast_children"]:
                self.render(child)
            self.footnote_paragraph = None

    def write_references(self):
        self.document.add_page_break()
        self.write_heading(Tr("References"), 1)
        lookup = self.main.references_viewer.reference_lookup
        for refid, entries in sorted(self.referenced.items()):
            reference = lookup[refid]
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(reference["id"])
            run.bold = True
            paragraph.add_run("  ")
            self.write_reference_authors(paragraph, reference)
            try:
                method = getattr(self, f"write_reference_{reference['type']}")
            except AttributeError:
                ic("unknown", reference["type"])
            else:
                method(paragraph, reference)
            self.write_reference_external_links(paragraph, reference)
            self.write_reference_xrefs(paragraph, entries)

    def write_reference_authors(self, paragraph, reference):
        count = len(reference["authors"])
        for pos, author in enumerate(reference["authors"]):
            if pos > 0:
                if pos == count - 1:
                    paragraph.add_run(" & ")
                else:
                    paragraph.add_run(", ")
            paragraph.add_run(utils.shortname(author))

    def write_reference_article(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}) ")
        try:
            paragraph.add_run(reference["title"].strip(".") + ". ")
        except KeyError:
            pass
        try:
            run = paragraph.add_run(f"{reference['journal']} ")
            run.font.italic = True
        except KeyError:
            pass
        try:
            paragraph.add_run(f"{reference['volume']} ")
        except KeyError:
            pass
        else:
            try:
                paragraph.add_run(f"({reference['number']})")
            except KeyError:
                pass
        try:
            paragraph.add_run(f": pp. {reference['pages'].replace('--', '-')}.")
        except KeyError:
            pass

    def write_reference_book(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}). ")
        run = paragraph.add_run(reference["title"].strip(".") + ". ")
        run.font.italic = True
        try:
            paragraph.add_run(f"{reference['publisher']}. ")
        except KeyError:
            pass

    def write_reference_link(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}). ")
        try:
            paragraph.add_run(reference["title"].strip(".") + ". ")
        except KeyError:
            pass
        try:
            add_hyperlink(paragraph, reference["url"], reference["title"])
        except KeyError:
            pass
        try:
            paragraph.add_run(f" Accessed {reference['accessed']}.")
        except KeyError:
            pass

    def write_reference_external_links(self, paragraph, reference):
        any_item = False
        for key, label, template in constants.REFERENCE_LINKS:
            try:
                value = reference[key]
                if any_item:
                    paragraph.add_run(", ")
                else:
                    paragraph.add_run("  ")
                add_hyperlink(
                    paragraph, template.format(value=value), f"{label}:{value}"
                )
                any_item = True
            except KeyError:
                pass

    def write_reference_xrefs(self, paragraph, entries):
        run = paragraph.add_run()
        run.add_break(docx.enum.text.WD_BREAK.LINE)
        paragraph.add_run("\t")
        entries.sort(key=lambda e: e["ordinal"])
        for entry in entries:
            paragraph.add_run(entry["heading"])
            if entry is not entries[-1]:
                paragraph.add_run(", ")

    def write_indexed(self):
        self.document.add_page_break()
        self.write_heading(Tr("Index"), 1)
        items = sorted(self.indexed.items(), key=lambda i: i[0].lower())
        for canonical, entries in items:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(canonical)
            run.bold = True
            paragraph.add_run("  ")
            entries.sort(key=lambda e: e["ordinal"])
            for entry in entries:
                paragraph.add_run(entry["heading"])
                if entry is not entries[-1]:
                    paragraph.add_run(", ")

    def render(self, ast):
        try:
            method = getattr(self, f"render_{ast['element']}")
        except AttributeError:
            ic("Could not handle ast", ast)
        else:
            method(ast)

    def render_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.render(child)

    def render_paragraph(self, ast):
        if self.footnote_paragraph:
            self.paragraph = self.footnote_paragraph
        else:
            self.paragraph = self.document.add_paragraph()
        if self.list_stack:
            data = self.list_stack[-1]
            depth = min(3, data["depth"])  # Max list depth in predef styles.
            if data["first_paragraph"]:
                if data["ordered"]:
                    if depth == 1:
                        style = self.document.styles["List Number"]
                    else:
                        style = self.document.styles[f"List Number {depth}"]
                else:
                    if depth == 1:
                        style = self.document.styles["List Bullet"]
                    else:
                        style = self.document.styles[f"List Bullet {depth}"]
            else:
                if depth == 1:
                    style = self.document.styles["List Continue"]
                else:
                    style = self.document.styles[f"List Continue {depth}"]
            data["first_paragraph"] = False
            self.paragraph.style = style
        else:
            self.paragraph.style = self.style_stack[-1]
        for child in ast["children"]:
            self.render(child)

    def render_raw_text(self, ast):
        line = ast["children"]
        line = line.rstrip("\n")
        run = self.paragraph.add_run(line)
        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True

    def render_blank_line(self, ast):
        pass

    def render_quote(self, ast):
        self.style_stack.append(QUOTE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_code_span(self, ast):
        run = self.paragraph.add_run(ast["children"])
        run.style = self.document.styles["Macro Text Char"]

    def render_code_block(self, ast):
        self.paragraph = self.document.add_paragraph(style=CODE_STYLE)
        self.style_stack.append(CODE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_fenced_code(self, ast):
        self.paragraph = self.document.add_paragraph(style=CODE_STYLE)
        self.style_stack.append(CODE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_emphasis(self, ast):
        self.italic = True
        for child in ast["children"]:
            self.render(child)
        self.italic = False

    def render_strong_emphasis(self, ast):
        self.bold = True
        for child in ast["children"]:
            self.render(child)
        self.bold = False

    def render_thematic_break(self, ast):
        paragraph = self.document.add_paragraph(constants.EM_DASH * 20)
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def render_link(self, ast):
        # XXX This handles only raw text within a link, nothing else.
        raw_text = []
        for child in ast["children"]:
            if child["element"] == "raw_text":
                raw_text.append(child["children"])
        add_hyperlink(self.paragraph, ast["dest"], "".join(raw_text))

    def render_list(self, ast):
        data = dict(
            ordered=ast["ordered"],
            bullet=ast["bullet"],  # Currently useless.
            start=ast["start"],  # Currently useless.
            tight=ast["tight"],  # Currently useless.
            count=0,  # Currently useless.
            depth=len(self.list_stack) + 1,
        )
        self.list_stack.append(data)
        for child in ast["children"]:
            self.render(child)
        self.list_stack.pop()

    def render_list_item(self, ast):
        data = self.list_stack[-1]
        data["count"] += 1  # Currently useless.
        data["first_paragraph"] = True
        for child in ast["children"]:
            self.render(child)

    def render_indexed(self, ast):
        entries = self.indexed.setdefault(ast["canonical"], [])
        self.indexed_count += 1
        entries.append(
            dict(
                id=f"i{self.indexed_count}",
                ordinal=self.current_text.ordinal,
                fullname=self.current_text.fullname,
                heading=self.current_text.heading,
            )
        )
        run = self.paragraph.add_run(ast["term"])
        font = self.config["indexed_font"]
        if font == constants.ITALIC:
            run.italic = True
        elif font == constants.BOLD:
            run.bold = True
        elif font == constants.UNDERLINE:
            run.underline = True

    def render_footnote_ref(self, ast):
        entries = self.footnotes.setdefault(self.current_text.fullname, {})
        label = int(ast["label"])
        entries[label] = dict(label=label)
        run = self.paragraph.add_run(str(label))
        run.font.superscript = True
        run.font.bold = True

    def render_footnote_def(self, ast):
        label = int(ast["label"])
        self.footnotes[self.current_text.fullname][label]["ast_children"] = ast[
            "children"
        ]

    def render_reference(self, ast):
        entries = self.referenced.setdefault(ast["reference"], [])
        self.referenced_count += 1
        entries.append(
            dict(
                ordinal=self.current_text.ordinal,
                fullname=self.current_text.fullname,
                heading=self.current_text.heading,
            )
        )
        run = self.paragraph.add_run(ast["reference"])
        font = self.config["references_font"]
        if font == constants.ITALIC:
            run.italic = True
        elif font == constants.BOLD:
            run.bold = True
        elif font == constants.UNDERLINE:
            run.underline = True


class Dialog(tk.simpledialog.Dialog):
    "Dialog to confirm or modify configuration before export."

    def __init__(self, master, source, config):
        self.source = source
        self.config = copy.deepcopy(config)
        self.result = None
        super().__init__(master, title=Tr("DOCX export"))

    def body(self, body):
        row = 0

        row += 1
        label = tk.ttk.Label(body, text=Tr("Directory"))
        label.grid(row=row, column=0, padx=4, sticky=tk.NE)
        self.dirpath_entry = tk.ttk.Entry(body, width=40)
        self.dirpath_entry.insert(0, self.config.get("dirpath") or os.getcwd())
        self.dirpath_entry.grid(row=row, column=1, sticky=tk.W)
        button = tk.ttk.Button(body, text=Tr("Choose"), command=self.change_dirpath)
        button.grid(row=row, column=3)

        row += 1
        label = tk.ttk.Label(body, text=Tr("Filename"))
        label.grid(row=row, column=0, padx=4, sticky=tk.NE)
        self.filename_entry = tk.ttk.Entry(body, width=40)
        self.filename_entry.insert(0, self.config.get("filename") or "book.docx")
        self.filename_entry.grid(row=row, column=1, sticky=tk.W)

        row += 1
        label = tk.ttk.Label(body, text=Tr("Page break level"))
        label.grid(row=row, column=0, padx=4, sticky=tk.NE)
        self.page_break_level_var = tk.IntVar(
            value=min(self.config.get("page_break_level", 1), self.source.max_level)
        )
        frame = tk.ttk.Frame(body)
        frame.grid(row=row, column=1, padx=4, sticky=tk.W)
        for level in range(1, self.source.max_level + 1):
            button = tk.ttk.Radiobutton(
                frame, text=str(level), variable=self.page_break_level_var, value=level
            )
            button.pack(anchor=tk.W)

        row += 1
        label = tk.ttk.Label(body, text=Tr("Indexing font"))
        label.grid(row=row, column=0, padx=4, sticky=tk.NE)
        self.indexed_font_var = tk.StringVar(
            value=self.config.get("indexed_font", constants.NORMAL)
        )
        frame = tk.ttk.Frame(body)
        frame.grid(row=row, column=1, padx=4, sticky=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Normal"),
            variable=self.indexed_font_var,
            value=constants.NORMAL,
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Italic"),
            variable=self.indexed_font_var,
            value=constants.ITALIC,
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame, text=Tr("Bold"), variable=self.indexed_font_var, value=constants.BOLD
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Underline"),
            variable=self.indexed_font_var,
            value=constants.UNDERLINE,
        )
        button.pack(anchor=tk.W)

        row += 1
        label = tk.ttk.Label(body, text=Tr("Reference font"))
        label.grid(row=row, column=0, padx=4, sticky=tk.NE)
        self.references_font_var = tk.StringVar(
            value=self.config.get("references_font", constants.NORMAL)
        )
        frame = tk.ttk.Frame(body)
        frame.grid(row=row, column=1, padx=4, sticky=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Normal"),
            variable=self.references_font_var,
            value=constants.NORMAL,
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Italic"),
            variable=self.references_font_var,
            value=constants.ITALIC,
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Bold"),
            variable=self.references_font_var,
            value=constants.BOLD,
        )
        button.pack(anchor=tk.W)
        button = tk.ttk.Radiobutton(
            frame,
            text=Tr("Underline"),
            variable=self.references_font_var,
            value=constants.UNDERLINE,
        )
        button.pack(anchor=tk.W)

    def apply(self):
        self.config["dirpath"] = self.dirpath_entry.get().strip() or os.getcwd()
        filename = self.filename_entry.get().strip() or constants.BOOK
        self.config["filename"] = os.path.splitext(filename)[0] + ".docx"
        self.config["page_break_level"] = self.page_break_level_var.get()
        self.config["indexed_font"] = self.indexed_font_var.get()
        self.config["references_font"] = self.references_font_var.get()
        self.result = self.config

    def change_dirpath(self):
        dirpath = tk.filedialog.askdirectory(
            parent=self,
            title=Tr("Directory"),
            initialdir=self.config.get("dirpath") or os.getcwd(),
            mustexist=True,
        )
        if dirpath:
            self.dirpath_entry.delete(0, tk.END)
            self.dirpath_entry.insert(0, dirpath)


# From https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
def add_hyperlink(paragraph, url, text, color="2222FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value.
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values.
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element.
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element.
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given.
    if not color is None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested.
    # XXX Does not seem to work? /Per Kraulis
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element.
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
