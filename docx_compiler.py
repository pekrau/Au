"DOCX compiler."

from icecream import ic

import os.path

import docx

import constants
import utils


class Compiler:
    "DOCX compiler."

    def __init__(self, source):
        self.source = source

    def write(self, filepath=None):
        if filepath is None:
            filepath = os.path.join(self.source.abspath, self.source.name + ".docx")
        self.document = docx.Document()
        self.document.add_heading(self.source.name, 0)
        for item in self.source.items:
            if item.is_section:
                self.write_section(item, level=0)
            else:
                self.write_text(item, level=0)
        self.document.save(filepath)

    def write_section(self, section, level):
        if level <= constants.DOCX_PAGEBREAK_LEVEL:
            self.document.add_page_break()
        self.document.add_heading(section.name, level)
        for item in section.items:
            if item.is_section:
                self.write_section(item, level=level+1)
            else:
                self.write_text(item, level=level+1)

    def write_text(self, text, level):
        if level < constants.DOCX_PAGEBREAK_LEVEL:
            self.document.add_page_break()
        self.document.add_heading(text.name, level)
        self.render(text.ast)

    def render(self, ast):
        try:
            method = getattr(self, f"render_{ast['element']}")
        except AttributeError:
            ic("Could not handle ast", ast)
        else:
            method(ast)

    def render_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.render(child)

    def render_paragraph(self, ast):
        self.paragraph = self.document.add_paragraph()
        for child in ast["children"]:
            self.render(child)

    def render_raw_text(self, ast):
        line = ast["children"]
        if line[-1] == "\n":
            line[-1] = " "
        self.paragraph.add_run(line)
